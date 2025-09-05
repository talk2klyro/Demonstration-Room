from docx import Document
from bs4 import BeautifulSoup
from typing import Tuple
from ..utils import write_temp_bytes

def html_to_docx_bytes(html: str) -> Tuple[bytes, str]:
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    # simple mapping
    for element in soup.body.children:
        if element.name is None:
            text = element.strip()
            if text:
                doc.add_paragraph(text)
            continue
        if element.name in ["h1", "h2", "h3"]:
            p = doc.add_heading(level=1 if element.name == "h1" else 2)
            p.add_run(element.get_text())
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == "pre":
            doc.add_paragraph(element.get_text(), style='Intense Quote')
        elif element.name == "table":
            # naive table conversion
            rows = element.find_all("tr")
            if not rows:
                continue
            cols = rows[0].find_all(["td", "th"])
            table = doc.add_table(rows=0, cols=len(cols))
            for r in rows:
                cells = r.find_all(["td", "th"])
                row_cells = table.add_row().cells
                for i, c in enumerate(cells):
                    row_cells[i].text = c.get_text()
        else:
            # fallback
            doc.add_paragraph(element.get_text())

    import io
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    data = bio.read()
    _, fname = write_temp_bytes(data, suffix=".docx")
    return data, fname
